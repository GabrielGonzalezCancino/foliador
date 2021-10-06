import os
import docx
import io
import sys
import win32print
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc =docx.Document('boceto.docx')



def numero_to_letras(numero):
    indicador = [("",""),("MIL","MIL"),("MILLON","MILLONES"),("MIL","MIL"),("BILLON","BILLONES")]
    entero = int(numero)
    decimal = int(round((numero - entero)*100))
	#print 'decimal : ',decimal 
    contador = 0
    numero_letras = ""
    while entero >0:
        a = entero % 1000
        if contador == 0:
            en_letras = convierte_cifra(a,1).strip()
        else :
            en_letras = convierte_cifra(a,0).strip()
        if a==0:
            numero_letras = en_letras+" "+numero_letras
        elif a==1:
            if contador in (1,3):
                numero_letras = indicador[contador][0]+" "+numero_letras
            else:
                numero_letras = en_letras+" "+indicador[contador][0]+" "+numero_letras
        else:
            numero_letras = en_letras+" "+indicador[contador][1]+" "+numero_letras
        numero_letras = numero_letras.strip()
        contador = contador + 1
        entero = int(entero / 1000)
    numero_letras = numero_letras
    #print ('numero: ',numero)
    print (numero_letras)
    
def convierte_cifra(numero,sw):
    lista_centana = ["",("CIEN","CIENTO"),"DOSCIENTOS","TRESCIENTOS","CUATROCIENTOS","QUINIENTOS","SEISCIENTOS","SETECIENTOS","OCHOCIENTOS","NOVECIENTOS"]
    lista_decena = ["",("DIEZ","ONCE","DOCE","TRECE","CATORCE","QUINCE","DIECISEIS","DIECISIETE","DIECIOCHO","DIECINUEVE"),("VEINTE","VEINTI"),("TREINTA","TREINTA Y "),("CUARENTA" , "CUARENTA Y "),("CINCUENTA" , "CINCUENTA Y "),("SESENTA" , "SESENTA Y "),("SETENTA" , "SETENTA Y "),("OCHENTA" , "OCHENTA Y "),("NOVENTA" , "NOVENTA Y ")]
    lista_unidad = ["",("UN" , "UNO"),"DOS","TRES","CUATRO","CINCO","SEIS","SIETE","OCHO","NUEVE"]
    centena = int (numero / 100)
    decena = int((numero -(centena * 100))/10)
    unidad = int(numero - (centena * 100 + decena * 10))
	#print "centena: ",centena, "decena: ",decena,'unidad: ',unidad
 
    texto_centena = ""
    texto_decena = ""
    texto_unidad = ""
	#Validad las centenas
    texto_centena = lista_centana[centena]
    if centena == 1:
        if (decena + unidad)!=0:
            texto_centena = texto_centena[1]
        else :
            texto_centena = texto_centena[0]
 
	#Valida las decenas
    texto_decena = lista_decena[decena]
    if decena == 1:
        texto_decena = texto_decena[unidad]
    if decena > 1:
        if unidad != 0 :
            texto_decena = texto_decena[1]
        else:
            texto_decena = texto_decena[0]
 	#Validar las unidades
 	#print "texto_unidad: ",texto_unidad
    if decena != 1:
        texto_unidad = lista_unidad[unidad]
        if unidad == 1:
            texto_unidad = texto_unidad[sw]
    return "%s %s%s" %(texto_centena,texto_decena,texto_unidad)



#---------- interfaz grafica------------------
raiz = Tk()


raiz.title("Foliador de notaria")
#raiz.geometry("650x350")
#raiz.resizable(False,False)
raiz.config(bg="blue")
#raiz.attributes('-alpha', 0.7)
#miframe=Frame(width="650" , height="350")
miFrame= Frame()

miFrame.pack()
miFrame.config(bg="#f4f1f1")
miFrame.config(bd=20)
miFrame.config(width="670" , height="300")
miFrame.config(relief="groove")




#labels
Label(miFrame, text="Pagina Inicial",bg="#f4f1f1", font=("Comic Sans MS",12)).place(x=30, y=30)
Label(miFrame, text="Pagina Final",bg="#f4f1f1", font=("Comic Sans MS",12)).place(x=370, y=30)

#imagen
miImagen=PhotoImage(file="logo,notaria.png")
Label(miFrame, image=miImagen ,bg="#f4f1f1").place(x=10, y=65)


#cuadro de texto

cuadroInicial=Entry(miFrame )
cuadroInicial.place(x=140, y=35)



cuadroFinal=Entry(miFrame)
cuadroFinal.place(x=465, y=35)

#boton de envio
def codigoBoton():
    salto=0
    posicion=0
    #inputs
    
    pinicial = int(cuadroInicial.get())

    pfinal = int(cuadroFinal.get()) 
    
    Label(miFrame, text="Procesado!", pady=20).place(x=330, y=220)
    
    if(pinicial <= pfinal):
        while  pinicial <= pfinal: 
    #converimos todo el print de la funcion para utilizarlo como variable   
        
            old_stdout = sys.stdout
            new_stdout = io.StringIO()
            sys.stdout = new_stdout
            numero_to_letras(pinicial)
            output = new_stdout.getvalue()
            sys.stdout = old_stdout
            print("Pagina ", pinicial,"...ok")
            salida=output[:-1]
            paragraph = doc.paragraphs[posicion]
            run = paragraph.add_run()
            run.add_picture('logito.png',width=docx.shared.Inches(0.3),height=docx.shared.Inches(0.3))
            run_2 = paragraph.add_run(" \t \t \t \t \t \t {}".format(salida))
        #doc.paragraphs[posicion].add_run("{}\t \t \t \t \t \t \t \t \t \t".format(salida)).add_picture('logito.png',width=docx.shared.Inches(0.3),height=docx.shared.Inches(0.3))    
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph("\t \t \t \t \t \t \t \t \t \t \t{}".format(pinicial))
            posicion= posicion + 2
            if(pinicial != pfinal):
                pinicial=pinicial + 1
                doc.add_page_break()
            else:
                pinicial=pinicial + 1
                break
    else:
        print("pagina inicial mas grande que pagina final ")

    

#--------------consola------------------
   

    
botonEnvio=Button(miFrame,text="Aceptar", command=codigoBoton,).place(x=315, y=200)
botonEnvio=Button(miFrame,text="Salir", command=raiz.destroy,).place(x=370, y=200)
raiz.mainloop()


doc.save('foliador.docx')
