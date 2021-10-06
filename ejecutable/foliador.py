import os
import docx
import io
import sys
import win32print
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH



doc =docx.Document('boceto.docx')


def numero_to_letras(numero):
    indicador = [("",""),("MIL","MIL"),("MILLON","MILLONES"),("MIL","MIL"),("BILLON","BILLONES")]
    entero = int(numero)
    decimal = int(round((numero - entero)*100))
	#print 'decimal : ',decimal 
    contador = 0
    numero_letras = ""
    while entero >0:
        a = entero % 1000
        if contador == 0:
            en_letras = convierte_cifra(a,1).strip()
        else :
            en_letras = convierte_cifra(a,0).strip()
        if a==0:
            numero_letras = en_letras+" "+numero_letras
        elif a==1:
            if contador in (1,3):
                numero_letras = indicador[contador][0]+" "+numero_letras
            else:
                numero_letras = en_letras+" "+indicador[contador][0]+" "+numero_letras
        else:
            numero_letras = en_letras+" "+indicador[contador][1]+" "+numero_letras
        numero_letras = numero_letras.strip()
        contador = contador + 1
        entero = int(entero / 1000)
    numero_letras = numero_letras
    #print ('numero: ',numero)
    print (numero_letras)
    
def convierte_cifra(numero,sw):
    lista_centana = ["",("CIEN","CIENTO"),"DOSCIENTOS","TRESCIENTOS","CUATROCIENTOS","QUINIENTOS","SEISCIENTOS","SETECIENTOS","OCHOCIENTOS","NOVECIENTOS"]
    lista_decena = ["",("DIEZ","ONCE","DOCE","TRECE","CATORCE","QUINCE","DIECISEIS","DIECISIETE","DIECIOCHO","DIECINUEVE"),("VEINTE","VEINTI"),("TREINTA","TREINTA Y "),("CUARENTA" , "CUARENTA Y "),("CINCUENTA" , "CINCUENTA Y "),("SESENTA" , "SESENTA Y "),("SETENTA" , "SETENTA Y "),("OCHENTA" , "OCHENTA Y "),("NOVENTA" , "NOVENTA Y ")]
    lista_unidad = ["",("UN" , "UNO"),"DOS","TRES","CUATRO","CINCO","SEIS","SIETE","OCHO","NUEVE"]
    centena = int (numero / 100)
    decena = int((numero -(centena * 100))/10)
    unidad = int(numero - (centena * 100 + decena * 10))
	#print "centena: ",centena, "decena: ",decena,'unidad: ',unidad
 
    texto_centena = ""
    texto_decena = ""
    texto_unidad = ""
	#Validad las centenas
    texto_centena = lista_centana[centena]
    if centena == 1:
        if (decena + unidad)!=0:
            texto_centena = texto_centena[1]
        else :
            texto_centena = texto_centena[0]
 
	#Valida las decenas
    texto_decena = lista_decena[decena]
    if decena == 1:
        texto_decena = texto_decena[unidad]
    if decena > 1:
        if unidad != 0 :
            texto_decena = texto_decena[1]
        else:
            texto_decena = texto_decena[0]
 	#Validar las unidades
 	#print "texto_unidad: ",texto_unidad
    if decena != 1:
        texto_unidad = lista_unidad[unidad]
        if unidad == 1:
            texto_unidad = texto_unidad[sw]
    return "%s %s%s" %(texto_centena,texto_decena,texto_unidad)

#pedir por consola numeros
print("Coloque pagina inicial")
pinicial= int(input())
print("Coloque pagina final")
pfinal= int(input())
salto=0
posicion=0

if(pinicial <= pfinal):
    while  pinicial <= pfinal: 
    #converimos todo el print de la funcion para utilizarlo como variable   
        
        old_stdout = sys.stdout
        new_stdout = io.StringIO()
        sys.stdout = new_stdout
        numero_to_letras(pinicial)
        output = new_stdout.getvalue()
        sys.stdout = old_stdout
        print("Pagina ", pinicial,"...ok")
        salida=output[:-1]
        paragraph = doc.paragraphs[posicion]
        run = paragraph.add_run()
        run.add_picture('logito.png',width=docx.shared.Inches(0.3),height=docx.shared.Inches(0.3))
        run_2 = paragraph.add_run(" \t \t \t \t \t \t {}".format(salida))
        #doc.paragraphs[posicion].add_run("{}\t \t \t \t \t \t \t \t \t \t".format(salida)).add_picture('logito.png',width=docx.shared.Inches(0.3),height=docx.shared.Inches(0.3))    
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("\t \t \t \t \t \t \t \t \t \t \t{}".format(pinicial))
        posicion= posicion + 2
        if(pinicial != pfinal):
            pinicial=pinicial + 1
            doc.add_page_break()
        else:
            pinicial=pinicial + 1
            break
        
    
       
else:
    print("pagina inicial mas grande que pagina final ")

    
doc.save('foliador.docx')


#imprime word creado

import win32api
import win32print
from glob import glob

# A List containing the system printers
all_printers = [printer[2] for printer in win32print.EnumPrinters(2)]
# Ask the user to select a printer
printer_num = int(input("Escoge una impresora:\n"+"\n".join([f"{n} {p}" for n, p in enumerate(all_printers)])+"\n"))
# set the default printer
win32print.SetDefaultPrinter(all_printers[printer_num])
#pdf_dir = "D:/path/to/pdf_dir/**/*"
pdf_dir = "C:/Users/Gabriel/Desktop/codigos/foliador/executable/foliador.docx"
for f in glob(pdf_dir, recursive=True):
    win32api.ShellExecute(0, "print", f, None,  ".",  0)

input("pulsa cualquier boton para salir")

from os import remove
remove("foliador.docx")